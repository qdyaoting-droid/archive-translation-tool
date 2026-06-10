"""
app.py - 国际联盟档案翻译工具
运行方式：streamlit run app.py
"""

import streamlit as st
import tempfile
import shutil
import subprocess
import time
import csv
import re
from pathlib import Path
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# ============================================================
# 页面设置
# ============================================================
st.set_page_config(
    page_title="国际联盟档案翻译工具",
    page_icon="📜",
    layout="wide"
)

# ============================================================
# 样式
# ============================================================
st.markdown("""
<style>
.main-title {
    font-size: 2rem;
    font-weight: bold;
    color: #1F4E79;
    margin-bottom: 0.5rem;
}
.sub-title {
    font-size: 1rem;
    color: #555;
    margin-bottom: 2rem;
}
.step-box {
    background: #F0F7FF;
    border-left: 4px solid #2E75B6;
    padding: 1rem;
    border-radius: 0 8px 8px 0;
    margin: 0.5rem 0;
}
.success-box {
    background: #F0FFF4;
    border-left: 4px solid #38A169;
    padding: 1rem;
    border-radius: 0 8px 8px 0;
}
.warning-box {
    background: #FFFBF0;
    border-left: 4px solid #D69E2E;
    padding: 1rem;
    border-radius: 0 8px 8px 0;
}
</style>
""", unsafe_allow_html=True)

# ============================================================
# 标题
# ============================================================
st.markdown('<div class="main-title">📜 国际联盟档案翻译工具</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">League of Nations Archive Translation Tool — 自动提取、分类、翻译并生成Word文档</div>', unsafe_allow_html=True)
st.divider()

# ============================================================
# 侧边栏：设置
# ============================================================
with st.sidebar:
    st.header("⚙️ 设置")

    api_key = st.text_input(
        "DeepSeek API Key",
        type="password",
        placeholder="sk-xxxxxxxxxxxxxxxx",
        help="在 platform.deepseek.com 获取"
    )

    st.divider()
    st.markdown("**翻译模型**")
    model = st.selectbox(
        "选择模型",
        ["deepseek-chat", "deepseek-reasoner"],
        index=0
    )

    st.divider()
    st.markdown("**关于**")
    st.markdown("""
    本工具专为1920-1940年代  
    国际联盟档案处理设计。
    
    支持英文、法文档案的：
    - 自动文件识别与分割
    - 中英法对照翻译
    - 学术引文自动生成
    - Word文档导出
    """)

# ============================================================
# 核心处理函数
# ============================================================

def extract_text_from_pdf(pdf_path, chunks_dir):
    """从PDF提取文字"""
    import fitz
    doc = fitz.open(pdf_path)
    chunks_dir = Path(chunks_dir)
    chunks_dir.mkdir(exist_ok=True)

    for i, page in enumerate(doc, start=1):
        text = page.get_text()
        out_file = chunks_dir / f"page_{i:03}.txt"
        out_file.write_text(
            f"===== PAGE {i} =====\n\n{text}",
            encoding="utf-8"
        )
    return len(doc)

def classify_pages(chunks_dir):
    """规则分类"""
    chunks_dir = Path(chunks_dir)
    results = []
    prev_classification = None

    for file in sorted(chunks_dir.glob("page_*.txt")):
        text = file.read_text(encoding="utf-8", errors="ignore")
        upper = text.upper()
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        classification = "CONTINUATION"

        if len(text.strip()) < 30:
            classification = "BLANK PAGE"
        elif any(kw in upper for kw in [
            "CLASSEMENT", "DOSSIER NO", "REGISTRY",
            "REMETTRE CE DOCUMENT", "LISTE DES PIÈCES", "SCHEDULE NO"
        ]):
            classification = "COVER PAGE"
        elif any(kw in upper for kw in [
            "DEAR SIR", "DEAR MR", "CHER MONSIEUR", "CHER SIR",
            "MONSIEUR LE SECRÉTAIRE", "YOUR EXCELLENCY",
            "MEMORANDUM", "NOTE BY THE SECRETARY", "NOTE DU SECRÉTAIRE",
            "NOTE BY SIR", "NOTE DE SIR", "REPORT OF THE COMMISSION",
            "RAPPORT DE LA COMMISSION", "RESOLUTION", "CIRCULAR",
            "TELEGRAM", "TÉLÉGRAMME", "DOCUMENTS DISPOSAL SHEET",
        ]):
            classification = "NEW DOCUMENT"
        elif any(kw in upper for kw in [
            "FIRST MEETING", "PREMIÈRE SÉANCE",
            "MINUTES OF THE", "PROCÈS-VERBAUX DE LA",
            "PERMANENT MANDATES COMMISSION",
            "COMMISSION PERMANENTE DES MANDATS",
        ]):
            is_first = any(kw in upper for kw in [
                "FIRST MEETING", "PREMIÈRE SÉANCE",
                "OPENING SPEECH", "LIST OF MEMBERS",
                "MINUTES OF THE", "PROCÈS-VERBAUX DE LA",
            ])
            if prev_classification in ["NEW DOCUMENT", "CONTINUATION"] and not is_first:
                classification = "CONTINUATION"
            else:
                classification = "NEW DOCUMENT"
        elif any(kw in upper for kw in ["ANNEX", "ANNEXE", "APPENDIX"]):
            first_line = lines[0].upper() if lines else ""
            if any(kw in first_line for kw in ["ANNEX", "ANNEXE", "APPENDIX"]):
                classification = "NEW DOCUMENT"

        prev_classification = classification
        results.append([file.name, classification])

    return results

def group_documents(classification_results):
    """分组文件"""
    groups = []
    current = None

    for page, cls in classification_results:
        if cls == "BLANK PAGE":
            continue
        if cls in ["COVER PAGE", "NEW DOCUMENT"]:
            if current:
                groups.append(current)
            current = {
                "document_id": len(groups) + 1,
                "start_page": page,
                "end_page": page,
                "start_classification": cls,
                "pages": [page],
            }
        else:
            if current:
                current["end_page"] = page
                current["pages"].append(page)

    if current:
        groups.append(current)

    return groups

def clean_ocr_text(text):
    """合并OCR换行碎片"""
    lines = text.split("\n")
    merged = []
    buffer = ""

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                merged.append(buffer.strip())
                buffer = ""
            merged.append("")
            continue
        if stripped.startswith("===== PAGE"):
            if buffer:
                merged.append(buffer.strip())
                buffer = ""
            merged.append(stripped)
            continue
        if buffer:
            if buffer.endswith("-"):
                buffer = buffer[:-1] + stripped
            elif buffer[-1] in ".!?:\"'»":
                merged.append(buffer.strip())
                buffer = stripped
            elif len(buffer) < 40 and stripped and stripped[0].isupper():
                merged.append(buffer.strip())
                buffer = stripped
            else:
                buffer += " " + stripped
        else:
            buffer = stripped

    if buffer:
        merged.append(buffer.strip())

    return "\n".join(merged)

SYSTEM_PROMPT = """
你是专业的历史档案翻译助手，专门处理1920-1940年代国际联盟（League of Nations）档案。
档案语言为英文、法文或两者混合。

【输出格式】

每份文件开头必须输出：
---
【档案信息】
文献标题（Title）：
文献形成时间（Date）：
发文者（From）：
收文者（To）：
文件编号（Document Number）：
原始页码（Original Page Number）：
档案馆藏所（Repository）：United Nations Library & Archives, Geneva
档案引文（Citation）：[生成完整英文学术引文，可直接用于论文脚注]
---

【档案引文规范】
格式："文件标题." 机构名, 档案编号, United Nations Library & Archives, Geneva, pp. 页码.

示例：
- 会议记录："Minutes of the Seventh Session Held at Geneva from 19th to 30th October 1925." Permanent Mandates Commission, C.648.M.237.1925.VI, United Nations Library & Archives, Geneva, pp. 152–154.
- 信件：Robert de Caix to Eric Drummond, May 16, 1926, File 1/38257/248, United Nations Library & Archives, Geneva.
- 备忘录："Note on Liquor Traffic." Sir Frederick D. Lugard, October 14, 1927, C.P.M.637, United Nations Library & Archives, Geneva.

【正文翻译格式】
每个完整段落为一组（同一段落合并，不要每行单独成段）：

Original:
[完整段落原文]

【中译】
[对应中文译文]

【版式保留】
- 【页眉】机密等级、机构名称
- 【印章】如 RECEIVED IN REGISTRY
- 【手写批注】
- 【签名】

无法辨认处标注：[原文模糊，无法辨认]

【译者注】首次出现专有名词时加注：
【译者注】League of Nations，即国际联盟（1920—1946）。

【研究重点】涉及以下主题在段落前加 ★：
- 酒类贸易管制、禁酒政策
- 托管制度、殖民治理
- 财政利益、税率
- 国际合作谈判

【已知人物】
- V. Catastini：国际联盟托管科官员
- Sir Frederick D. Lugard：常设托管委员会英国委员
- William Rappard：托管科主任
- Robert de Caix：法国外交部亚洲司官员
- Sir Eric Drummond：国际联盟秘书长
"""

def translate_document(client, model, text, pages):
    """翻译单份文件"""
    cleaned = clean_ocr_text(text)
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"请按规范翻译以下档案。\n原始页码：{pages}\n\n{cleaned}"}
        ],
        max_tokens=8192,
        temperature=0.1
    )
    return response.choices[0].message.content

# ============================================================
# Word生成函数
# ============================================================

def set_font(run, size=10.5, bold=False, italic=False,
             color=None, en_font="Times New Roman", zh_font="SimSun"):
    run.font.name = en_font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), zh_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def set_line_spacing(para, spacing_pt=20):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(int(spacing_pt * 20)))
    spacing.set(qn("w:lineRule"), "exact")
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)

def set_shading(para, fill_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    existing = pPr.find(qn("w:shd"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)

def add_divider(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_word(translations):
    """生成Word文档，返回bytes"""
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 封面
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("国际联盟档案翻译数据库")
    set_font(title_run, size=20, bold=True, color=(0x1F, 0x4E, 0x79))

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("League of Nations Archive Translation Database")
    set_font(sub_run, size=12, color=(0x55, 0x55, 0x55))

    doc.add_paragraph()
    add_divider(doc)

    for i, (doc_id, text) in enumerate(translations, 1):
        # 文件标题
        h_para = doc.add_paragraph()
        h_run = h_para.add_run(f"文件 {doc_id:03}")
        set_font(h_run, size=14, bold=True, color=(0x1F, 0x4E, 0x79))
        add_divider(doc)

        # 解析翻译内容
        lines = text.split("\n")
        j = 0
        while j < len(lines):
            line = lines[j].strip()
            if not line or line == "---":
                j += 1
                continue

            if line == "【档案信息】":
                info_para = doc.add_paragraph()
                info_run = info_para.add_run("档案信息")
                set_font(info_run, size=11, bold=True, color=(0x2E, 0x75, 0xB6))
                j += 1
                continue

            if line.startswith("档案引文（Citation）：") or line.startswith("Citation:"):
                ref_text = line.split("：", 1)[-1].strip() if "：" in line else line.split(":", 1)[-1].strip()
                if ref_text:
                    # 引文标签
                    label_para = doc.add_paragraph()
                    label_para.paragraph_format.left_indent = Cm(0.5)
                    label_run = label_para.add_run("📎 档案引文（Citation）")
                    set_font(label_run, size=9, bold=True, color=(0x40, 0x40, 0x40))
                    # 引文内容
                    ref_para = doc.add_paragraph()
                    ref_para.paragraph_format.left_indent = Cm(1)
                    ref_para.paragraph_format.space_after = Pt(8)
                    set_shading(ref_para, "EBEBEB")
                    ref_run = ref_para.add_run(ref_text)
                    set_font(ref_run, size=10, italic=True, color=(0x20, 0x20, 0x20))
                j += 1
                continue

            if any(line.startswith(k) for k in [
                "文献标题", "文献形成时间", "发文者", "收文者",
                "文件编号", "原始页码", "档案馆藏所"
            ]):
                info_para = doc.add_paragraph()
                info_para.paragraph_format.left_indent = Cm(0.5)
                set_line_spacing(info_para, 18)
                info_run = info_para.add_run(line)
                set_font(info_run, size=10)
                j += 1
                continue

            if line.startswith("【译者注】"):
                note_para = doc.add_paragraph()
                note_para.paragraph_format.left_indent = Cm(0.8)
                set_line_spacing(note_para, 18)
                note_run = note_para.add_run(line)
                set_font(note_run, size=9, italic=True, color=(0xC0, 0x00, 0x00))
                j += 1
                continue

            if line.startswith("★"):
                hl_para = doc.add_paragraph()
                hl_para.paragraph_format.left_indent = Cm(0.8)
                set_shading(hl_para, "FFF2CC")
                set_line_spacing(hl_para, 20)
                hl_run = hl_para.add_run(line)
                set_font(hl_run, size=10, bold=True)
                j += 1
                continue

            if any(line.startswith(k) for k in ["【页眉】", "【页脚】", "【印章】", "【手写批注】", "【签名】"]):
                stamp_para = doc.add_paragraph()
                stamp_para.paragraph_format.left_indent = Cm(0.8)
                set_line_spacing(stamp_para, 18)
                stamp_run = stamp_para.add_run(line)
                set_font(stamp_run, size=9, color=(0x70, 0x70, 0x70))
                j += 1
                continue

            if line.startswith("Original:"):
                orig_text = line[9:].strip()
                j += 1
                while j < len(lines):
                    nl = lines[j].strip()
                    if not nl or nl.startswith("【") or nl.startswith("Original:") or nl.startswith("★"):
                        break
                    orig_text += " " + nl
                    j += 1
                if orig_text:
                    orig_para = doc.add_paragraph()
                    orig_para.paragraph_format.left_indent = Cm(0.8)
                    orig_para.paragraph_format.space_after = Pt(2)
                    set_line_spacing(orig_para, 20)
                    label = orig_para.add_run("Original: ")
                    set_font(label, size=10, bold=True, color=(0x55, 0x55, 0x55))
                    content = orig_para.add_run(orig_text)
                    set_font(content, size=10)
                continue

            if line.startswith("【中译】"):
                trans_text = line[5:].strip()
                j += 1
                while j < len(lines):
                    nl = lines[j].strip()
                    if not nl or nl.startswith("【") or nl.startswith("Original:") or nl.startswith("★"):
                        break
                    trans_text += " " + nl
                    j += 1
                if trans_text:
                    trans_para = doc.add_paragraph()
                    trans_para.paragraph_format.left_indent = Cm(0.8)
                    trans_para.paragraph_format.space_after = Pt(8)
                    set_line_spacing(trans_para, 20)
                    set_shading(trans_para, "EBF3FB")
                    label = trans_para.add_run("【中译】")
                    set_font(label, size=10, bold=True, color=(0x2E, 0x75, 0xB6))
                    content = trans_para.add_run(trans_text)
                    set_font(content, size=10)
                continue

            # 普通行
            normal_para = doc.add_paragraph()
            normal_para.paragraph_format.left_indent = Cm(0.5)
            set_line_spacing(normal_para, 20)
            normal_run = normal_para.add_run(line)
            set_font(normal_run, size=10)
            j += 1

        if i < len(translations):
            doc.add_page_break()

    # 转为bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ============================================================
# 主界面
# ============================================================

col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("📁 上传档案PDF")
    uploaded_files = st.file_uploader(
        "支持同时上传多份PDF",
        type=["pdf"],
        accept_multiple_files=True,
        help="每份PDF会单独处理，生成对应的Word文件"
    )

with col2:
    st.subheader("📊 状态")
    if not api_key:
        st.warning("请在左侧填入 DeepSeek API Key")
    elif not uploaded_files:
        st.info("请上传PDF文件")
    else:
        st.success(f"已上传 {len(uploaded_files)} 份PDF，可以开始处理")

st.divider()

# 开始处理按钮
if uploaded_files and api_key:
    if st.button("🚀 开始翻译", type="primary", use_container_width=True):

        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

        # 检查pymupdf
        try:
            import fitz
        except ImportError:
            st.error("请先安装 PyMuPDF：pip3 install pymupdf")
            st.stop()

        results = {}  # {文件名: word_bytes}

        for pdf_file in uploaded_files:
            st.markdown(f"---")
            st.markdown(f"### 📄 处理：{pdf_file.name}")

            with tempfile.TemporaryDirectory() as tmpdir:
                tmpdir = Path(tmpdir)
                chunks_dir = tmpdir / "chunks"
                chunks_dir.mkdir()

                # 保存PDF
                pdf_path = tmpdir / pdf_file.name
                pdf_path.write_bytes(pdf_file.read())

                # 步骤1：提取文字
                with st.status("📖 正在提取文字...", expanded=False) as status:
                    try:
                        page_count = extract_text_from_pdf(str(pdf_path), str(chunks_dir))
                        status.update(label=f"✅ 提取完成，共 {page_count} 页", state="complete")
                    except Exception as e:
                        status.update(label=f"❌ 提取失败：{e}", state="error")
                        continue

                # 步骤2：分类
                with st.status("🔍 正在识别文件结构...", expanded=False) as status:
                    try:
                        classification = classify_pages(str(chunks_dir))
                        groups = group_documents(classification)
                        status.update(label=f"✅ 识别完成，共 {len(groups)} 份独立文件", state="complete")
                    except Exception as e:
                        status.update(label=f"❌ 分类失败：{e}", state="error")
                        continue

                # 步骤3：翻译
                translations = []
                progress_bar = st.progress(0, text="正在翻译...")

                for idx, group in enumerate(groups):
                    pages = group["pages"]
                    progress_bar.progress(
                        (idx + 1) / len(groups),
                        text=f"正在翻译第 {idx+1}/{len(groups)} 份文件..."
                    )

                    # 合并页面文字
                    full_text = ""
                    for page in pages:
                        page_file = chunks_dir / page
                        if page_file.exists():
                            full_text += page_file.read_text(encoding="utf-8") + "\n\n"

                    if len(full_text.strip()) < 50:
                        continue

                    try:
                        result = translate_document(
                            client, model, full_text,
                            ", ".join(pages)
                        )
                        translations.append((group["document_id"], result))
                        time.sleep(0.5)
                    except Exception as e:
                        st.warning(f"文件 {idx+1} 翻译失败：{e}")

                progress_bar.progress(1.0, text="翻译完成！")

                # 步骤4：生成Word
                with st.status("📝 正在生成Word文档...", expanded=False) as status:
                    try:
                        word_bytes = generate_word(translations)
                        output_name = Path(pdf_file.name).stem + "_translated.docx"
                        results[output_name] = word_bytes
                        status.update(label="✅ Word文档生成完成", state="complete")
                    except Exception as e:
                        status.update(label=f"❌ Word生成失败：{e}", state="error")
                        continue

                # 下载按钮
                st.markdown(
                    f'<div class="success-box">✅ <b>{pdf_file.name}</b> 处理完成！'
                    f'共翻译 {len(translations)} 份文件。</div>',
                    unsafe_allow_html=True
                )
                st.download_button(
                    label=f"⬇️ 下载 {output_name}",
                    data=word_bytes,
                    file_name=output_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        st.balloons()
        st.success("🎉 所有PDF处理完成！")

elif uploaded_files and not api_key:
    st.warning("请先在左侧边栏填入 DeepSeek API Key")
