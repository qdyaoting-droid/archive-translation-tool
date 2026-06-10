"""
generate_word.py (改进版)
改进：
1. 字体设置更稳定
2. 段落格式更清晰
3. 引文格式高亮显示
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 路径配置
# ============================================================
TRANSLATIONS_DIR = Path("output/translations")
OUTPUT_FILE      = Path("output/final_archive.docx")

# ============================================================
# 字体设置（稳定版）
# ============================================================
def set_font(run, size=10.5, bold=False, italic=False,
             color=None, en_font="Times New Roman", zh_font="SimSun"):
    run.font.name = en_font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 设置中文字体
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), zh_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def set_line_spacing(para, spacing_pt=20):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(int(spacing_pt * 20)))
    spacing.set(qn("w:lineRule"), "exact")
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)

def set_shading(para, fill_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    existing = pPr.find(qn("w:shd"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)

def add_divider(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

# ============================================================
# 段落类型
# ============================================================
def add_heading1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    set_font(run, size=14, bold=True, color=(0x1F, 0x4E, 0x79))
    return para

def add_heading2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_font(run, size=11, bold=True, color=(0x2E, 0x75, 0xB6))
    return para

def add_archive_info(doc, text):
    """档案信息行"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    set_line_spacing(para, 18)
    run = para.add_run(text)
    set_font(run, size=10)
    return para

def add_reference(doc, text):
    """引文行（深灰底，学术引用格式）"""
    # 标题行
    title_para = doc.add_paragraph()
    title_para.paragraph_format.left_indent = Cm(0.5)
    title_para.paragraph_format.space_before = Pt(4)
    set_line_spacing(title_para, 18)
    title_run = title_para.add_run("📎 档案引文（Citation）")
    set_font(title_run, size=9, bold=True, color=(0x40, 0x40, 0x40))

    # 引文内容
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_after = Pt(8)
    set_line_spacing(para, 18)
    set_shading(para, "EBEBEB")
    content = para.add_run(text)
    set_font(content, size=10, italic=True, color=(0x20, 0x20, 0x20))
    return para

def add_original(doc, text):
    """英法原文段落（仅用于引文/专有名词例外情况）"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    para.paragraph_format.space_after = Pt(2)
    set_line_spacing(para, 20)
    set_shading(para, "F5F5F5")
    label = para.add_run("原文：")
    set_font(label, size=9, bold=True, color=(0x55, 0x55, 0x55))
    content = para.add_run(text)
    set_font(content, size=9, italic=True)
    return para

def add_translation(doc, text):
    """中文译文段落（正文，无底色标签）"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    set_line_spacing(para, 20)
    content = para.add_run(text)
    set_font(content, size=10.5)
    return para

def add_page_marker(doc, text):
    """页码标记"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_font(run, size=9, color=(0x70, 0x70, 0x70))
    return para

def add_right_aligned(doc, text):
    """靠右对齐段落（落款、签名等）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_after = Pt(4)
    set_line_spacing(para, 20)
    run = para.add_run(text)
    set_font(run, size=10.5)
    return para

def add_note(doc, text):
    """译者注（红色）"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    set_line_spacing(para, 18)
    run = para.add_run(f"【译者注】{text}")
    set_font(run, size=9, italic=True, color=(0xC0, 0x00, 0x00))
    return para

def add_highlight(doc, text):
    """研究重点（黄底）"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    set_line_spacing(para, 20)
    set_shading(para, "FFF2CC")
    run = para.add_run(f"★ {text}")
    set_font(run, size=10, bold=True)
    return para

def add_stamp(doc, text):
    """印章/页眉/批注"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    set_line_spacing(para, 18)
    run = para.add_run(text)
    set_font(run, size=9, color=(0x70, 0x70, 0x70))
    return para

def add_normal(doc, text):
    """普通行"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    set_line_spacing(para, 20)
    run = para.add_run(text)
    set_font(run, size=10)
    return para

# ============================================================
# 解析翻译文本
# ============================================================
def parse_and_write(doc, text):
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line or line == "---":
            i += 1
            continue

        # 档案信息块标题
        if line == "【档案信息】":
            add_heading2(doc, "档案信息")
            i += 1
            continue

        # 引文行
        if line.startswith("档案引文（Citation）：") or line.startswith("引文格式（Reference）：") or line.startswith("Citation:"):
            ref_text = ""
            if "：" in line:
                ref_text = line.split("：", 1)[-1].strip()
            elif ":" in line:
                ref_text = line.split(":", 1)[-1].strip()
            # 引文可能跨多行（直到空行或下一个字段）
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or any(next_line.startswith(k) for k in ["文献", "发文者", "收文者", "文件编号", "档案", "---", "【"]):
                    break
                ref_text += " " + next_line
                i += 1
            if ref_text:
                add_reference(doc, ref_text.strip())
            continue

        # 档案信息各字段
        if any(line.startswith(k) for k in [
            "文献标题", "文献形成时间", "发文者", "收文者",
            "文件编号", "原始页码", "档案馆藏所"
        ]):
            add_archive_info(doc, line)
            i += 1
            continue

        # 译者注
        if line.startswith("【译者注】"):
            add_note(doc, line[6:])
            i += 1
            continue

        # 研究重点
        if line.startswith("★"):
            add_highlight(doc, line[1:].strip())
            i += 1
            continue

        # 页码标记 【第X页】【PDF第X页 / 文件第Y页】
        if line.startswith("【第") and ("页】" in line or "页 /" in line):
            add_page_marker(doc, line)
            i += 1
            continue

        # 印章/页眉/批注
        if any(line.startswith(k) for k in [
            "【页眉】", "【页脚】", "【印章】",
            "【手写批注】", "【分发名单】"
        ]):
            add_stamp(doc, line)
            i += 1
            continue

        # 签名/落款（靠右对齐）
        if line.startswith("【签名】"):
            add_right_aligned(doc, line[4:].strip() or line)
            i += 1
            continue

        # Original 块（兼容旧格式，仅用于例外情况原文保留）
        if line.startswith("Original:"):
            original_text = line[9:].strip()
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or next_line.startswith("【") or next_line.startswith("Original:") or next_line.startswith("★"):
                    break
                original_text += " " + next_line
                i += 1
            if original_text:
                add_original(doc, original_text)
            continue

        # 【中译】块（兼容旧格式）
        if line.startswith("【中译】"):
            trans_text = line[5:].strip()
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or next_line.startswith("【") or next_line.startswith("Original:") or next_line.startswith("★"):
                    break
                trans_text += " " + next_line
                i += 1
            if trans_text:
                add_translation(doc, trans_text)
            continue

        # Translation: 格式（兼容旧格式）
        if line.startswith("Translation:"):
            trans_text = line[12:].strip()
            if trans_text:
                add_translation(doc, trans_text)
            i += 1
            continue

        # 普通正文行（新格式下直接是译文）
        add_translation(doc, line)
        i += 1

# ============================================================
# 主流程
# ============================================================
def main():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 封面
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(60)
    title_run = title_para.add_run("国际联盟档案翻译数据库")
    set_font(title_run, size=20, bold=True, color=(0x1F, 0x4E, 0x79))

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("League of Nations Archive Translation Database")
    set_font(sub_run, size=12, color=(0x55, 0x55, 0x55))

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    # 读取翻译文件
    txt_files = sorted(TRANSLATIONS_DIR.glob("doc_*.txt"))

    if not txt_files:
        print("❌ 没有找到翻译文件，请先运行 translate_documents.py")
        return

    print(f"共找到 {len(txt_files)} 份翻译文件，正在生成 Word...")

    for txt_file in txt_files:
        doc_id = txt_file.stem.upper()

        # 文件标题
        add_heading1(doc, f"文件 {doc_id}")
        add_divider(doc)

        text = txt_file.read_text(encoding="utf-8")
        parse_and_write(doc, text)

        doc.add_page_break()

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\n✅ Word 文档已生成 → {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
